"""
FinGuard FINAL_REPORT → Beautiful DOCX Generator
=================================================
Converts the markdown report into a professional, classic Word document
with proper styling, embedded diagrams, formatted tables, and clean typography.
"""

import os
import re
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DIAGRAMS = os.path.join(BASE_DIR, "docs", "diagrams")
OUT_PATH = os.path.join(BASE_DIR, "FINAL_REPORT_CLEANED.docx")


def setup_styles(doc):
    """Configure document-wide styles for a classic, professional look."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Georgia'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15

    for level in range(1, 4):
        sname = f'Heading {level}'
        hstyle = doc.styles[sname]
        hstyle.font.name = 'Calibri'
        hstyle.font.color.rgb = RGBColor(0x0D, 0x0D, 0x0D)
        hstyle.font.bold = True
        if level == 1:
            hstyle.font.size = Pt(22)
            hstyle.paragraph_format.space_before = Pt(24)
            hstyle.paragraph_format.space_after = Pt(12)
        elif level == 2:
            hstyle.font.size = Pt(16)
            hstyle.paragraph_format.space_before = Pt(18)
            hstyle.paragraph_format.space_after = Pt(8)
        else:
            hstyle.font.size = Pt(13)
            hstyle.paragraph_format.space_before = Pt(12)
            hstyle.paragraph_format.space_after = Pt(6)


def set_cell_shading(cell, color_hex):
    """Set background shading for a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_table(doc, headers, rows, caption=None, header_color="1A1A1A"):
    """Create a beautifully formatted table."""
    num_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cell, header_color)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg = "F9F9F9" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            # Handle bold markers
            val_str = str(val)
            if val_str.startswith('**') and val_str.endswith('**'):
                run = p.add_run(val_str.strip('*'))
                run.bold = True
            else:
                run = p.add_run(val_str)
            run.font.size = Pt(9)
            run.font.name = 'Georgia'
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_shading(cell, bg)

    # Set column widths evenly
    for row in table.rows:
        for cell in row.cells:
            cell.width = Inches(6.5 / num_cols)

    if caption:
        p = doc.add_paragraph()
        run = p.add_run(caption)
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(12)

    return table


def add_image_safe(doc, filename, width=Inches(5.8), caption=None):
    """Add an image if it exists, with a caption."""
    path = os.path.join(DIAGRAMS, filename)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=width)
        if caption:
            cp = doc.add_paragraph()
            run = cp.add_run(caption)
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.space_after = Pt(10)
        return True
    return False


def add_para(doc, text, bold=False, italic=False, size=Pt(11), color=None, align=None, space_after=Pt(6)):
    """Add a styled paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Georgia'
    run.font.size = size
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if align:
        p.alignment = align
    p.paragraph_format.space_after = space_after
    return p


def add_rich_para(doc, parts, align=None, space_after=Pt(6)):
    """Add a paragraph with mixed formatting. parts = [(text, bold, italic), ...]"""
    p = doc.add_paragraph()
    for text, bold, italic in parts:
        run = p.add_run(text)
        run.font.name = 'Georgia'
        run.font.size = Pt(11)
        run.bold = bold
        run.italic = italic
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    if align:
        p.alignment = align
    p.paragraph_format.space_after = space_after
    return p


def add_bullet(doc, text, level=0, bold_prefix=None):
    """Add a bullet point, optionally with a bold prefix."""
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.name = 'Georgia'
        run.font.size = Pt(11)
        run = p.add_run(text)
        run.font.name = 'Georgia'
        run.font.size = Pt(11)
    else:
        p.text = ''
        run = p.add_run(text)
        run.font.name = 'Georgia'
        run.font.size = Pt(11)
    return p


def add_numbered(doc, text, bold_prefix=None):
    """Add a numbered list item."""
    p = doc.add_paragraph(style='List Number')
    p.text = ''
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.name = 'Georgia'
        run.font.size = Pt(11)
        run = p.add_run(text)
        run.font.name = 'Georgia'
        run.font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.font.name = 'Georgia'
        run.font.size = Pt(11)
    return p


def add_code_block(doc, code_text):
    """Add a code block with monospace font and gray background."""
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    # Gray background via shading
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2"/>')
    p._p.get_or_add_pPr().append(shading)
    return p


def add_formula(doc, formula_text):
    """Add a math formula as styled text."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(formula_text)
    run.font.name = 'Cambria Math'
    run.font.size = Pt(11)
    run.italic = True
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    return p


def add_divider(doc):
    """Add a thin horizontal line as a divider."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="CCCCCC"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    return p


# ═══════════════════════════════════════════════════════════════════════
# BUILD THE DOCUMENT
# ═══════════════════════════════════════════════════════════════════════
def build_docx():
    doc = Document()
    setup_styles(doc)

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── COVER PAGE ──────────────────────────────────────────────────
    doc.add_paragraph()  # spacer
    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('CYBER SECURITY INNOVATION CHALLENGE 1.0')
    run.font.name = 'Calibri'
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = RGBColor(0x0D, 0x0D, 0x0D)

    add_para(doc, '', size=Pt(6))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('PROTOTYPE DEVELOPMENT: Stage III')
    run.font.name = 'Calibri'
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    add_divider(doc)

    for line in [
        'Problem Statement Domain: Mule Accounts & Collusive Fraud in UPI',
        'Problem Statement: Detection of Mule Account Networks and Coordinated Fund Laundering in UPI Payment Ecosystems',
    ]:
        add_para(doc, line, size=Pt(11), align=WD_ALIGN_PARAGRAPH.CENTER,
                 color=RGBColor(0x44, 0x44, 0x44))

    add_para(doc, '', size=Pt(12))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('FinGuard')
    run.font.name = 'Calibri'
    run.font.size = Pt(36)
    run.bold = True
    run.font.color.rgb = RGBColor(0x0D, 0x0D, 0x0D)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Real-Time Multi-Signal Mule Account Detection System')
    run.font.name = 'Georgia'
    run.font.size = Pt(14)
    run.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    add_para(doc, '', size=Pt(24))

    # Project Links table on cover
    add_styled_table(doc,
        ['Resource', 'Link'],
        [
            ['GitHub Repository', '[Insert GitHub Repo URL]'],
            ['Live Deployed URL', '[Insert Deployed URL]'],
            ['Dashboard Demo Video (YouTube)', '[Insert YouTube Link]'],
            ['Pitch Deck / Presentation', '[Insert Link if applicable]'],
        ],
        header_color="333333"
    )

    add_para(doc, '', size=Pt(12))

    # Team table on cover
    add_styled_table(doc,
        ['Role', 'Name', 'Institute', 'Enrolment No.', 'Email ID'],
        [
            ['Team Lead', '[Name]', '[Institute]', '[Enrolment No.]', '[Email]'],
            ['Member 2', '[Name]', '[Institute]', '[Enrolment No.]', '[Email]'],
            ['Member 3', '[Name]', '[Institute]', '[Enrolment No.]', '[Email]'],
            ['Member 4', '[Name]', '[Institute]', '[Enrolment No.]', '[Email]'],
            ['Member 5', '[Name]', '[Institute]', '[Enrolment No.]', '[Email]'],
        ],
        header_color="333333"
    )

    doc.add_page_break()

    # ── TABLE OF CONTENTS ───────────────────────────────────────────
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. Project Overview and Team Details',
        '2. Problem Statement and Background',
        '3. Literature Review / Existing Solutions',
        '4. Proposed Solution and Technical Architecture',
        '5. Innovation and Novelty Elements',
        '6. Unique Selling Proposition (USP), Business Model, and Industry Relevance',
        '7. Prototype Demonstration, Security Testing, and Deployment Details',
        '8. Limitations and Challenges',
        '9. Roadmap Towards MVP',
        '10. Team Composition and Individual Contributions',
        '11. References',
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.name = 'Georgia'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        p.paragraph_format.space_after = Pt(3)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # SECTION 1
    # ══════════════════════════════════════════════════════════════════
    doc.add_heading('1. Project Overview and Team Details', level=1)

    doc.add_heading('1.1 Project Summary', level=2)
    add_para(doc, "FinGuard detects mule accounts in India's UPI ecosystem, and it does so in real time. Mule accounts, for anyone unfamiliar, are bank accounts that criminals use as pass-throughs for laundering stolen money. They're a massive headache for the payments industry because each individual transaction through a mule looks perfectly normal. It's only when you zoom out and look at the bigger picture (the network structure, the timing, the devices involved) that the fraud becomes visible.")

    add_image_safe(doc, '01_system_flow.png', Inches(5.5),
                   'Figure 1: FinGuard high-level system flow, from transaction ingestion to dashboard output.')

    add_para(doc, "Our system tackles this with what we call a five-signal ensemble. We pull together behavioral profiling, transaction graph analysis, device fingerprinting, temporal pattern detection, and unsupervised ML into a single scoring pipeline. The whole idea is that even if a sophisticated mule operator manages to fool one of those signals, the other four will still catch them.")

    doc.add_heading('1.2 Objective', level=2)
    add_para(doc, 'We set out to build a working prototype with four concrete goals:')
    add_numbered(doc, ' mule accounts by crunching transaction patterns, network topology, device correlations, and timing anomalies, all at once rather than one at a time.', bold_prefix='Detect')
    add_numbered(doc, ' every account into a risk tier (CRITICAL, HIGH, MEDIUM, or LOW) and back that up with specific, readable evidence. No black-box scores.', bold_prefix='Classify')
    add_numbered(doc, ' the fraud networks on an interactive dashboard so that human investigators can actually make sense of what\'s going on and take action.', bold_prefix='Visualize')
    add_numbered(doc, ' Sub-50ms scoring per account. UPI doesn\'t wait around, and neither can we.', bold_prefix='Stay fast.')

    doc.add_heading('1.3 Scope', level=2)
    add_para(doc, 'The prototype covers the full pipeline end-to-end:')
    add_bullet(doc, ' We use synthetic but realistic UPI transaction data covering six different fraud scenarios: star aggregation, circular networks, chain laundering, device rings, rapid onboarding fraud, and night-time smurfing.', bold_prefix='Data Ingestion:')
    add_bullet(doc, ' Five independent scoring modules feed into a weighted ensemble with confidence boosting when multiple signals agree.', bold_prefix='Detection Engine:')
    add_bullet(doc, ' 11 endpoints with API-key auth, rate limiting (120 req/min), structured audit logs, and telemetry baked in.', bold_prefix='REST API:')
    add_bullet(doc, ' An 8-tab React SPA covering command center, risk analysis, ML insights, network graph, timeline, alerts, live API testing, and an about page.', bold_prefix='Dashboard:')
    add_bullet(doc, ' Docker containers with multi-service orchestration. We run as non-root, have health checks, the works.', bold_prefix='Deployment:')

    doc.add_heading('1.4 Team Details', level=2)
    add_para(doc, 'Team Name: FinGuard', bold=True)
    add_para(doc, 'Our team of five split the work roughly along these lines:')
    add_styled_table(doc,
        ['Member', 'Role', 'Institute', 'Key Contribution'],
        [
            ['[Team Lead Name]', 'Team Lead', '[Institute]', 'System architecture, risk engine design, project coordination'],
            ['[Member 2]', 'Backend Developer', '[Institute]', 'Detection modules, API development, ML pipeline'],
            ['[Member 3]', 'Frontend Developer', '[Institute]', 'React dashboard, data visualization, UX design'],
            ['[Member 4]', 'Data Engineer', '[Institute]', 'Data generation, testing, validation scenarios'],
            ['[Member 5]', 'DevOps / Documentation', '[Institute]', 'Docker deployment, security hardening, documentation'],
        ],
    )
    add_para(doc, "In practice, the boundaries were fuzzy; everybody ended up debugging graph algorithms at 2 AM at some point.", italic=True, color=RGBColor(0x55, 0x55, 0x55))

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # SECTION 2
    # ══════════════════════════════════════════════════════════════════
    doc.add_heading('2. Problem Statement and Background', level=1)

    doc.add_heading('2.1 Context: The UPI Ecosystem', level=2)
    add_para(doc, "UPI is enormous. In October 2024 alone, it processed 13.1 billion transactions worth \u20b920.64 lakh crore [1]. That makes it the largest real-time payment system anywhere in the world, not just in India. The whole design philosophy behind UPI is speed and accessibility: instant transfers, near-zero cost, interoperability between banks. It's genuinely transformative for financial inclusion.")
    add_para(doc, "But that same openness creates a problem. When you build a system that lets anyone move money to anyone instantly, you've also built something that's extremely attractive to criminals.")

    doc.add_heading('2.2 The Mule Account Problem', level=2)
    add_para(doc, "So what exactly is a mule account? Simple: it's a bank account used as a waystation for dirty money. Someone gets defrauded through phishing or a vishing call, their money goes into a mule account, and from there it gets shuffled around through more accounts until it's cashed out somewhere untraceable. The person whose account is being used as a mule may or may not be aware. Sometimes they're recruited via fake job postings, sometimes they're completely in the dark.")
    add_para(doc, "Here's what makes mule detection so frustrating:")
    add_numbered(doc, " A \u20b95,000 UPI transfer from one valid account holder to another? Nothing suspicious about that on its own.", bold_prefix="Individual transactions look fine.")
    add_numbered(doc, " Mule operations create distinctive network shapes (star patterns, sequential chains, circular loops) but you can't see any of that if you're only looking at one transaction at a time.", bold_prefix="The fraud is in the coordination.")
    add_numbered(doc, " Mule accounts often show burst activity followed by total silence, or they operate mostly between midnight and 5 AM. The signatures are there, but they're easy to miss.", bold_prefix="Timing gives them away, but subtly.")
    add_numbered(doc, " When five different bank accounts are all being operated from the same phone, that's not a coincidence. But traditional per-account monitoring won't catch this.", bold_prefix="Device sharing is a huge tell.")

    doc.add_heading('2.3 Limitations of Current Approaches', level=2)
    add_para(doc, "Most fraud detection systems running inside UPI today are basically rule engines. They check each transaction against a list of thresholds. These work, up to a point. But they have some pretty fundamental blind spots:")
    add_bullet(doc, " Each transaction gets evaluated in isolation. A rule engine has no concept of an account receiving money from five accounts and immediately forwarding it all.", bold_prefix="They can't see networks.")
    add_bullet(doc, " Mule operators adapt fast. By the time a fraud team manually writes a new rule, the criminals have moved on.", bold_prefix="Rules go stale.")
    add_bullet(doc, " You either block the transaction or you don't. No nuanced risk scoring.", bold_prefix="It's all-or-nothing.")
    add_bullet(doc, " A rule engine doesn't know how to combine weird transaction timing AND shared devices AND star-shaped graph topology.", bold_prefix="No cross-signal thinking.")
    add_para(doc, "The RBI clearly recognises how serious this has gotten. They've mandated enhanced fraud monitoring under the Digital Payments Security Controls directions [2], and NPCI's own guidelines now push for real-time, multi-dimensional detection capabilities [3].")

    doc.add_heading('2.4 Problem Formulation', level=2)
    add_para(doc, "Mathematically, here's what we're trying to do. Take a set of UPI accounts A = {a\u2081, a\u2082, \u2026, a\u2099} and build a transaction graph G = (A, E) where each edge represents a fund transfer. We want to compute a risk score R(a\u1d62) \u2208 [0, 100] for every account:")
    add_formula(doc, 'R(a\u1d62) = \u03a3\u2096 w\u2096 \u00b7 S\u2096(a\u1d62) + Boost({S\u2096(a\u1d62)})')
    add_para(doc, "The five S\u2096 terms correspond to our five detection signals (behavioral, graph, device, temporal, ML). Each has a weight w\u2096. The Boost function is a confidence amplifier that rewards agreement across signals.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # SECTION 3
    # ══════════════════════════════════════════════════════════════════
    doc.add_heading('3. Literature Review / Existing Solutions', level=1)

    doc.add_heading('3.1 Rule-Based Systems', level=2)
    add_para(doc, "The bread-and-butter of UPI fraud detection today is still rule engines sitting inside the payment switch. They're fast, explainable, and for catching obvious fraud they work fine. But coordinated multi-account fraud is a blind spot. A rule engine doesn't cross-reference accounts against each other [4].")

    doc.add_heading('3.2 Machine Learning Approaches', level=2)
    add_para(doc, "People have thrown all the usual ML suspects at banking fraud: Random Forests, Gradient Boosting, deep neural nets [5]. In the specific context of mule accounts, they run into trouble:")
    add_bullet(doc, " Mule accounts are almost never explicitly labelled in production data.", bold_prefix="Where are the labels?")
    add_bullet(doc, " Mule tactics evolve constantly. A model trained on last quarter's patterns may be useless by next month.", bold_prefix="Concept drift hits hard.")
    add_bullet(doc, " Most ML models consume a feature vector per account. They miss the relational network structure.", bold_prefix="They look at accounts individually.")

    doc.add_heading('3.3 Graph-Based Detection', level=2)
    add_para(doc, "Graph approaches are where the recent literature gets genuinely exciting:")
    add_bullet(doc, "Jambhrunkar et al. (2025) built MuleTrack, a lightweight temporal learning framework for money mule detection [6].")
    add_bullet(doc, "Cheng et al. (2024) showed that GNN-based methods consistently beat traditional feature-based classifiers for coordinated fraud [7].")
    add_bullet(doc, "Caglayan and Bahtiyar (2022) used Node2Vec embeddings for money laundering detection with improved detection rates [8].")
    add_bullet(doc, "Huang (2025) used community detection for finding money mules on transaction graphs [9].")
    add_bullet(doc, "Neo4j (2023) published industry case studies on graph databases for fraud detection [10].")

    doc.add_heading('3.4 Unsupervised Anomaly Detection', level=2)
    add_para(doc, "Isolation Forest [11] deserves its own mention. The core intuition is elegant: anomalies are few and different, so they can be isolated by random binary splits much more quickly than normal data points. You don't need labelled examples. Combine this with Z-score outlier detection, and you get an ensemble that can spot novel fraud patterns it's never been explicitly trained on.")

    doc.add_heading('3.5 Gap Analysis', level=2)
    add_para(doc, "We did a systematic comparison to figure out where the gaps are:")
    add_styled_table(doc,
        ['Capability', 'Rule-Based', 'Supervised ML', 'Graph Methods', 'FinGuard (Ours)'],
        [
            ['Temporal Pattern Detection', 'Static thresholds', 'Limited', 'None', 'Full (5 sub-signals)'],
            ['Network/Graph Analysis', 'None', 'None', 'Yes', 'Yes (3 patterns + DFS)'],
            ['Device Correlation', 'None', 'Partial', 'None', 'Full'],
            ['Unsupervised (No labels)', 'N/A', 'No', 'Partial', 'Yes (IF + Z-score)'],
            ['Real-Time Scoring', 'Fast', 'Moderate', 'Slow', 'Fast (<50ms)'],
            ['Explainability', 'Clear', 'Black-box', 'Limited', 'Full (3\u20135 evidence items)'],
            ['Multi-Signal Ensemble', 'No', 'No', 'No', 'Yes (5-factor weighted)'],
            ['Confidence Boosting', 'No', 'No', 'No', 'Yes (multi-signal)'],
        ],
        caption='Table 1: Comparison of FinGuard with existing detection approaches.'
    )
    add_para(doc, "The bottom line: nobody's combined all five signal types into one ensemble with explainable evidence and confidence-aware boosting. That's the gap we're filling.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # SECTION 4
    # ══════════════════════════════════════════════════════════════════
    doc.add_heading('4. Proposed Solution and Technical Architecture', level=1)

    doc.add_heading('4.1 Overall Architecture', level=2)
    add_para(doc, "We went with a layered architecture. Not because it's trendy, but because it genuinely makes sense here. Each layer does one thing, and doing it well. Data comes in at the bottom, gets processed through five parallel scoring modules, aggregated into a final risk score, served through an API, and displayed on a dashboard.")

    add_image_safe(doc, '02_architecture.png', Inches(5.8),
                   'Figure 2: System architecture for FinGuard Mule Account Detection Platform.')

    add_para(doc, "The key thing to notice is that the five detection modules all run independently. They don't talk to each other. That's deliberate. We want uncorrelated signals so the confidence boosting at the aggregation layer actually means something.")

    doc.add_heading('4.2 Detection Methodology', level=2)
    add_para(doc, "This is the heart of the system. We go through each module in detail.")

    doc.add_heading('4.2.1 Behavioral Analysis (Weight: 25%)', level=3)
    add_para(doc, "The behavioral module looks at each account's transaction patterns individually and asks: does this look like a mule? It checks six things:")
    add_numbered(doc, " How many transactions has this account made? 10+ gets a score of 35. 5+ gets 25.", bold_prefix="Velocity Detection:")
    add_numbered(doc, " We calculate the pass-through ratio (outflow / inflow). If between 0.8 and 1.2, the account is forwarding almost everything. +35.", bold_prefix="Asymmetric Flow Analysis:")
    add_numbered(doc, " Unusually high average amounts (over \u20b95,000) add +20. Single transactions above \u20b910,000 add +15.", bold_prefix="Amount Anomalies:")
    add_numbered(doc, " An account less than 7 days old with 2+ transactions scores +40. Mule operators frequently open fresh accounts.", bold_prefix="New Account + Rapid Activity:")
    add_numbered(doc, " Cumulative volume exceeds \u20b950,000 adds +20.", bold_prefix="Total Volume Spikes:")
    add_numbered(doc, " Accounts that only send money but never receive any get +20.", bold_prefix="Unidirectional Flow:")

    doc.add_heading('4.2.2 Graph Analytics (Weight: 40%)', level=3)
    add_para(doc, "This module gets the highest weight, and for good reason. Mule operations are fundamentally network problems. We build a directed graph G = (V, E) using NetworkX and hunt for three specific topologies:")

    add_image_safe(doc, '03_graph_patterns.png', Inches(5.8),
                   'Figure 3: The three mule network topologies detected by the graph analytics module.')

    add_rich_para(doc, [('Star Pattern (Aggregator/Distributor): ', True, False),
                        ('5+ inflows funnelling to 1 outflow: +45. 3+ inflows to 1 outflow: +30. 1 inflow splitting to 5+ outflows: +45.', False, False)])
    add_rich_para(doc, [('Chain Pattern (Layered Laundering): ', True, False),
                        ('A\u2192B\u2192C\u2192D\u2192E sequential chain. 6+ hops: +35, 5 hops: +30, 4 hops: +20. BFS with depth limits.', False, False)])
    add_rich_para(doc, [('Circular Pattern (Fund Rotation): ', True, False),
                        ('Money loops back to origin. Custom DFS cycle detector with depth cap of 6 (O(V\u00b7d) time). Cycle of length 3\u20136: +50.', False, False)])

    doc.add_heading('4.2.3 Device Fingerprinting (Weight: 15%)', level=3)
    add_para(doc, "If ten accounts are all logging in from the same device, something is very wrong. No legitimate user has ten bank accounts on one phone.")
    add_bullet(doc, " Account sharing a device with 10+ others: +50. With 5+: +40. With 3+: +30.", bold_prefix="Device Concentration:")
    add_bullet(doc, " Single account accessed from 5+ different devices: +30. From 3+: +20.", bold_prefix="Multi-Device Rotation:")

    add_image_safe(doc, '04_device_flow.png', Inches(5.5),
                   'Figure 4: Device fingerprinting detection flow and scoring thresholds.')

    doc.add_heading('4.2.4 Temporal Analysis (Weight: 10%)', level=3)
    add_numbered(doc, " 3+ transactions within 60 seconds: +35. Within 5 minutes: +25.", bold_prefix="Burst Detection:")
    add_numbered(doc, " Over half an account's transactions between midnight and 5 AM: +30.", bold_prefix="Odd-Hour Activity:")
    add_numbered(doc, " Second half of timeline has 3x more transactions than first half: +25.", bold_prefix="Velocity Spikes:")
    add_numbered(doc, " 70%+ transactions on weekends with at least 4 weekend transactions: +15.", bold_prefix="Weekend Concentration:")
    add_numbered(doc, " CV of time gaps below 0.15 with mean gap under 600 seconds (bot signature): +30.", bold_prefix="Uniform Timing:")

    doc.add_heading('4.2.5 ML Anomaly Detection (Weight: 10%)', level=3)
    add_para(doc, "We deliberately built our own Isolation Forest in pure NumPy (~200 lines). 100 trees, 256 samples each. 17 features per account covering transaction counts, amount statistics, network metrics, account metadata, device metrics, and derived metrics.")
    add_formula(doc, 'S_IF(x) = 2^(\u2212E[h(x)] / c(n))')
    add_para(doc, "The final ML score blends Isolation Forest and Z-score detection:")
    add_formula(doc, 'S_ML = 0.7 \u00b7 S_IF + 0.3 \u00b7 S_Z-score')
    add_para(doc, "We also implemented permutation-based feature importance and SHAP-like local explanations listing the top 5 features driving each account's anomaly score.")

    doc.add_heading('4.3 Risk Aggregation and Confidence Boosting', level=2)
    add_formula(doc, 'R_base = 0.25\u00b7S_B + 0.40\u00b7S_G + 0.15\u00b7S_D + 0.10\u00b7S_T + 0.10\u00b7S_ML')

    add_styled_table(doc,
        ['Active Signals', 'Boost'],
        [
            ['\u22654 signals above threshold', '+20'],
            ['\u22653 signals above threshold', '+15'],
            ['\u22652 signals above threshold', '+8'],
            ['Graph \u226530 AND Device \u226515', '+10'],
            ['Behavioral \u226530 AND Graph \u226530', '+8'],
            ['Behavioral \u226540 AND Graph \u226540 AND Device \u226530', '+12'],
        ],
        caption='Table 2: Multi-signal confidence boosting rules.'
    )

    add_styled_table(doc,
        ['Risk Level', 'Score Range', 'Recommended Action'],
        [
            ['CRITICAL', '\u226585', 'Block immediately, freeze account, file SAR'],
            ['HIGH', '70\u201384', 'Manual investigation within 24 hours'],
            ['MEDIUM', '40\u201369', 'Add to watchlist, periodic review'],
            ['LOW', '<40', 'Allow, routine monitoring'],
        ],
        caption='Table 3: Risk classification thresholds and recommended actions.'
    )

    add_image_safe(doc, '05_risk_pipeline.png', Inches(5.8),
                   'Figure 5: Complete risk scoring pipeline with weighted aggregation and confidence boosting.')

    doc.add_heading('4.4 API Layer', level=2)
    add_para(doc, "The backend runs on FastAPI v2.1.0 and exposes 11 endpoints:")
    add_styled_table(doc,
        ['Endpoint', 'Method', 'Function'],
        [
            ['/score/{account_id}', 'GET', 'Real-time single account scoring'],
            ['/batch_score', 'POST', 'Batch scoring with graph & ML caching'],
            ['/simulate', 'POST', 'Transaction simulation with dual-side risk'],
            ['/stats', 'GET', 'System-wide risk distribution statistics'],
            ['/api/dashboard', 'GET', 'Pre-computed dashboard data (all scores)'],
            ['/api/network', 'GET', 'Graph nodes/edges for vis-network rendering'],
            ['/api/timeline', 'GET', 'Transaction timeline and temporal heatmaps'],
            ['/api/report', 'GET', 'Auto-generated Markdown investigation report'],
            ['/metrics', 'GET', 'SRE/observability performance metrics'],
            ['/health', 'GET', 'Container health check endpoint'],
            ['/docs', 'GET', 'Interactive Swagger API documentation'],
        ],
        caption='Table 4: REST API endpoint reference.'
    )
    add_para(doc, "Security: API-key authentication (X-API-Key header), rate limiting at 120 requests/60-second window per IP, CORS locked to specific origins, structured JSON audit logging with unique request IDs.")

    doc.add_heading('4.5 Frontend Dashboard', level=2)
    add_para(doc, "The frontend is a React 18 + Vite single-page app with seven tabs:")
    add_numbered(doc, ' The landing page. Risk distribution, summary stats, signal heatmap, riskiest account cards.', bold_prefix='Command Center:')
    add_numbered(doc, ' Searchable, filterable table. Full forensic breakdown per account.', bold_prefix='Risk Analysis:')
    add_numbered(doc, ' Feature importance rankings and anomaly score distributions with per-account SHAP-like explanations.', bold_prefix='ML Insights:')
    add_numbered(doc, ' Interactive vis-network with risk-coded colours. Filter by risk level.', bold_prefix='Network Graph:')
    add_numbered(doc, ' Transaction timeline with hourly bucketing and day\u00d7hour heatmap.', bold_prefix='Timeline:')
    add_numbered(doc, ' Built-in API testing: score accounts, run simulations, batch scoring.', bold_prefix='Real-Time API:')
    add_numbered(doc, ' Documentation, methodology, architecture overview embedded in the app.', bold_prefix='About:')

    add_image_safe(doc, '10_signal_radar.png', Inches(5.5),
                   'Figure 6: Signal contribution radar comparing a mule account vs a legitimate account.')

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # SECTION 5
    # ══════════════════════════════════════════════════════════════════
    doc.add_heading('5. Innovation and Novelty Elements', level=1)

    innovations = [
        ('5.1 Five-Signal Ensemble Architecture', "Most fraud detection systems rely on one, maybe two signals. Nobody was combining all five (behavioral, graph, device, temporal, and ML) into one weighted ensemble. The advantage isn't just coverage; it's redundancy. A clever mule operator might game one signal, but evading all five simultaneously is vanishingly unlikely."),
        ('5.2 Multi-Signal Confidence Boosting', "When two signals both flag an account, that's more meaningful than either alone. Our boosting mechanism formalises this with additive bumps of +8 to +20. Because our modules don't share internals, agreement carries real evidential weight. Same logic as ensemble learning in ML, applied at the risk aggregation level."),
        ('5.3 Zero-Label ML Detection', "We wrote the Isolation Forest from scratch in NumPy. ~200 lines. No scikit-learn. It finds anomalies without labelled data. Deploy on a brand new platform on day one, with zero historical fraud labels, and it starts flagging outliers immediately."),
        ('5.4 Efficient Graph Algorithms', "We learned the hard way that nx.simple_cycles() has exponential worst-case complexity. Our custom DFS cycle detector with depth cap of 6 runs in O(V\u00b7d) time. Same for BFS chain detection. The practical difference for real-time scoring is night and day."),
        ('5.5 Explainability by Design', "Every detection module generates 3\u20135 specific evidence items as part of its core logic. Things like \"Star-pattern: 5 inflows \u2192 1 outflow\" or \"Burst: 4 transactions within 60 seconds.\" Investigators immediately understand why an account was flagged."),
        ('5.6 Production-Grade Security', "API-key auth, per-IP rate limiting (120 req/min), CORS restricted to specific origins, structured JSON audit logging with request IDs, non-root Docker execution. Table-stakes for anything touching production in a bank."),
    ]
    for title, text in innovations:
        doc.add_heading(title, level=2)
        add_para(doc, text)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # SECTION 6
    # ══════════════════════════════════════════════════════════════════
    doc.add_heading('6. Unique Selling Proposition (USP), Business Model, and Industry Relevance', level=1)

    doc.add_heading('6.1 Key Differentiators', level=2)
    add_rich_para(doc, [('1. Five signals where others use one or two. ', True, False),
                        ('Suppose a mule has behavioral score 30 and graph pattern 45. Neither alone crosses thresholds. But our ensemble with boosting pushes combined score to 68+. A single-model system would miss it.', False, False)])
    add_rich_para(doc, [('2. We put graphs first. ', True, False),
                        ("Graph analytics gets 40% weight. Mule operations are, at their core, network operations. Only graph analysis reveals the star pattern, the chain, the circular loop of an organized ring.", False, False)])
    add_rich_para(doc, [('3. No training data required. ', True, False),
                        ("Our Isolation Forest runs unsupervised. Deploy on a brand new platform, give it current data, and it starts flagging outliers right away.", False, False)])
    add_rich_para(doc, [('4. Full explainability. ', True, False),
                        ("Every score comes with component breakdowns, evidence items, confidence levels, and recommended actions. RBI regulations demand transparency in fraud monitoring.", False, False)])

    doc.add_heading('6.2 Industry Relevance', level=2)
    add_rich_para(doc, [('For Banks and PSPs: ', True, False),
                        ("Catching mule networks before cash-out saves real money. Pre-classified risk tiers and auto-generated evidence reduce investigation time. Helps meet RBI compliance requirements.", False, False)])
    add_rich_para(doc, [('For NPCI and UPI Infrastructure: ', True, False),
                        ("Network-wide view of mule operations spanning the entire ecosystem. Runs within UPI's latency constraints.", False, False)])
    add_rich_para(doc, [('For Law Enforcement: ', True, False),
                        ("Auto-generated investigation reports with network graphs, temporal heatmaps, and complete mule network identification. Forensic-grade audit logs.", False, False)])

    doc.add_heading('6.3 Competitive Positioning', level=2)
    add_styled_table(doc,
        ['Feature', 'Static Rules', 'Single ML Model', 'FinGuard'],
        [
            ['Detection Signals', '1 (rules)', '1 (features)', '5 (ensemble)'],
            ['Graph Awareness', 'None', 'None', 'Full'],
            ['Device Correlation', 'None', 'Partial', 'Full'],
            ['Labeled Data Required', 'No', 'Yes', 'No'],
            ['Explainability', 'High', 'Low', 'High'],
            ['Real-Time Capable', 'Yes', 'Moderate', 'Yes (<50ms)'],
            ['Confidence Levels', 'No', 'No', 'Yes'],
            ['Recommended Actions', 'Manual', 'None', 'Automated'],
            ['Deployment Complexity', 'Low', 'High', 'Medium (containerized)'],
        ],
        caption='Table 5: Competitive positioning of FinGuard against existing approaches.'
    )

    doc.add_heading('6.4 Business Model and Market Viability', level=2)
    add_rich_para(doc, [('Target Market: ', True, False),
                        ("India has 300+ banks and 50+ PSPs on UPI. Most still rely on static rule engines. RBI's push for enhanced fraud monitoring creates active regulatory pressure to upgrade. Primary targets: mid-tier banks and growing PSPs.", False, False)])
    add_rich_para(doc, [('Market Size: ', True, False),
                        ("UPI processes 13.1 billion transactions/month. Fraud losses estimated in thousands of crores annually. Each mule investigation costs banks \u20b915,000\u201325,000 in analyst time. Global fraud detection market projected to cross $60 billion by 2027.", False, False)])

    add_styled_table(doc,
        ['Tier', 'Target Customer', 'Pricing (Indicative)'],
        [
            ['Starter', 'Small fintechs, neobanks', '\u20b950,000/month or \u20b90.02/transaction'],
            ['Enterprise', 'Mid-tier banks, PSPs', '\u20b93\u20135 lakh/month with SLA'],
            ['Infrastructure', 'NPCI, large banks', 'Custom licensing, annual contracts'],
        ],
        caption='Revenue model tiers.'
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # SECTION 7
    # ══════════════════════════════════════════════════════════════════
    doc.add_heading('7. Prototype Demonstration, Security Testing, and Deployment Details', level=1)

    doc.add_heading('7.1 Prototype Overview', level=2)
    add_para(doc, "Everything described in Section 4 is implemented and working. Not a design doc, but a running system.")
    add_styled_table(doc,
        ['Component', 'Technology', 'Version'],
        [
            ['Backend API', 'FastAPI + Uvicorn', '2.1.0 / 0.27.1'],
            ['Frontend', 'React + Vite', '18.x / 5.x'],
            ['Language', 'Python', '3.11'],
            ['Graph Engine', 'NetworkX', '3.2.1'],
            ['ML Engine', 'Custom Isolation Forest (NumPy)', '1.26.4'],
            ['Data Processing', 'Pandas', '2.2.0'],
            ['Visualization', 'Plotly', '5.18.0'],
            ['Containerization', 'Docker + Compose', 'Multi-stage build'],
        ],
        caption='Table 6: Technology stack.'
    )

    doc.add_heading('7.2 Test Data and Scenarios', level=2)
    add_styled_table(doc,
        ['Scenario', 'Pattern', 'Accounts', 'Expected Risk'],
        [
            ['Star Aggregator', '5 sources \u2192 1 mule \u2192 1 dist. \u2192 3 sinks', '10', 'CRITICAL/HIGH'],
            ['Circular Network', '4-node loop (A\u2192B\u2192C\u2192D\u2192A), shared device', '4', 'CRITICAL'],
            ['Chain Laundering', '5-node sequential chain', '5', 'HIGH'],
            ['Device Ring', '3 accounts, 1 shared device', '3', 'HIGH/MEDIUM'],
            ['Rapid Onboarding', '1-day account, 8 receives + 5 sends in 30 min', '1+', 'CRITICAL'],
            ['Night Smurfing', '12+ transactions between 1\u20134 AM', '1+', 'HIGH'],
        ],
        caption='Table 7: Test scenarios with expected detection outcomes.'
    )
    add_para(doc, "We also seeded 25+ legitimate background accounts with normal transaction patterns to validate against false positives.")

    doc.add_heading('7.3 Demo Flow', level=2)
    add_para(doc, "Here's how a typical demo runs:")
    add_code_block(doc, 'docker-compose up --build\n# Backend: http://localhost:8000\n# Frontend: http://localhost:5173\n# API Docs: http://localhost:8000/docs')
    add_numbered(doc, " Dashboard loads with bird's-eye overview: accounts analysed, risk distribution, signal heatmap.", bold_prefix="Command Center:")
    add_numbered(doc, " Search by account ID, filter by risk level. Full forensic breakdown per account.", bold_prefix="Risk Analysis:")
    add_numbered(doc, " Interactive vis-network with colour-coded nodes (red=CRITICAL, orange=HIGH, yellow=MEDIUM, green=LOW).", bold_prefix="Network Graph:")
    add_numbered(doc, " Feature importance rankings, anomaly score distributions, per-account SHAP-like explanations.", bold_prefix="ML Insights:")
    add_numbered(doc, " Score accounts, simulate transactions, batch scoring directly from the dashboard.", bold_prefix="Live API Testing:")
    add_numbered(doc, " Auto-generated Markdown investigation report with executive summary and recommendations.", bold_prefix="Generate Report:")

    doc.add_heading('7.4 Detection Results', level=2)
    add_para(doc, "Every single mule scenario gets caught:")
    add_styled_table(doc,
        ['Scenario', 'Key Account', 'Score', 'Level', 'Primary Evidence'],
        [
            ['Star Aggregator', 'mule_aggregator@upi', '92', 'CRITICAL', 'Star pattern, 95% pass-through, 5-day account'],
            ['Circular Network', 'circle_node_1@upi', '88', 'CRITICAL', 'Circular network, shared device (4 accounts)'],
            ['Chain Laundering', 'chain_node_2@upi', '76', 'HIGH', 'Deep chain (4+ hops), high velocity'],
            ['Device Ring', 'device_ring_1@upi', '72', 'HIGH', 'Shared device (3 accounts), young account'],
            ['Rapid Onboarding', 'new_mule_account@upi', '95', 'CRITICAL', '1-day account, burst (8 txns/20 min)'],
            ['Night Smurfing', 'smurf_master@upi', '78', 'HIGH', 'Night activity (85%), burst pattern'],
        ],
        caption='Table 8: Detection results (all mule accounts correctly identified).'
    )

    add_image_safe(doc, '06_risk_distribution.png', Inches(5.5),
                   'Figure 7: Detection scores across all six mule scenarios and overall risk distribution.')

    add_para(doc, "Legitimate accounts: all scored below 40 (LOW). 0% false positives and 100% mule detection on this dataset.", bold=True)

    doc.add_heading('7.5 Performance Metrics', level=2)
    add_styled_table(doc,
        ['Metric', 'Value'],
        [
            ['Single account scoring latency', '<50ms'],
            ['Batch scoring (30 accounts)', '<500ms'],
            ['API startup time', '<3s'],
            ['Graph construction time', '<100ms'],
            ['ML model training + scoring', '<2s'],
            ['Memory footprint', '<150MB'],
        ],
        caption='Table 9: Performance benchmarks.'
    )

    doc.add_heading('7.6 Real-World Deployment Considerations', level=2)
    add_styled_table(doc,
        ['Prototype Component', 'Production Equivalent'],
        [
            ['CSV data files', 'Kafka stream ingestion \u2192 PostgreSQL'],
            ['NetworkX in-memory graph', 'Neo4j graph database'],
            ['In-memory caching', 'Redis cache layer'],
            ['Single Docker container', 'Kubernetes cluster with auto-scaling'],
            ['File-based audit logs', 'ELK stack (Elasticsearch, Logstash, Kibana)'],
            ['Pickle model store', 'MLflow model registry'],
        ],
    )
    add_image_safe(doc, '07_deployment.png', Inches(5.5),
                   'Figure 8: Prototype to production deployment migration path.')

    doc.add_heading('7.7 Security Testing Results', level=2)
    add_para(doc, "We ran an automated security test suite (scripts/security_test.py) against the live API. 42 test cases across 8 categories. All 42 passed.", bold=True)
    add_styled_table(doc,
        ['Category', 'Tests', 'Passed', 'Result'],
        [
            ['API Key Authentication', '6', '6', '100%'],
            ['Injection Attacks (SQL, NoSQL, XSS, Cmd, Path)', '10', '10', '100%'],
            ['Rate Limiting', '3', '3', '100%'],
            ['CORS Policy', '3', '3', '100%'],
            ['Input Validation & Error Handling', '7', '7', '100%'],
            ['HTTP Method Restriction', '4', '4', '100%'],
            ['Security Headers & Info Leakage', '5', '5', '100%'],
            ['Audit Logging', '4', '4', '100%'],
            ['Total', '42', '42', '100%'],
        ],
        caption='Table 10: Security test results by category.'
    )

    add_para(doc, "Key findings:", bold=True)
    add_bullet(doc, ' 10 injection payloads (SQL, NoSQL, XSS, command injection, path traversal, template injection) caused no crashes, no data leaks, no reflected XSS.', bold_prefix='Injection Resistance:')
    add_bullet(doc, ' Rate limiter triggered after 103 rapid requests, returned HTTP 429. Recovered after window reset.', bold_prefix='Rate Limiting:')
    add_bullet(doc, " Allowed origins accepted, malicious origins blocked, no wildcard (*) CORS configured.", bold_prefix='CORS Policy:')
    add_bullet(doc, ' Every response includes X-Request-Id and X-Response-Time. No stack traces or internal paths leaked.', bold_prefix='Security Headers:')
    add_bullet(doc, ' Structured JSON audit log with timestamps, event types, request IDs, methods, paths, IPs, status codes, and response times.', bold_prefix='Audit Trail:')

    doc.add_heading('OWASP API Security Top 10 Mapping', level=3)
    add_styled_table(doc,
        ['OWASP Risk', 'Our Mitigation', 'Status'],
        [
            ['API1: Broken Object Level Auth', 'Account IDs are lookup keys only', 'Addressed'],
            ['API2: Broken Authentication', 'API-key auth on protected endpoints', 'Addressed'],
            ['API3: Broken Property Level Auth', 'Pydantic strict schemas', 'Addressed'],
            ['API4: Unrestricted Resource Consumption', 'Rate limiting 120 req/min per IP', 'Addressed'],
            ['API5: Broken Function Level Auth', 'HTTP method restrictions per endpoint', 'Addressed'],
            ['API6: Unrestricted Sensitive Flows', 'Rate limiting + API key', 'Addressed'],
            ['API7: SSRF', 'No outbound HTTP calls', 'N/A'],
            ['API8: Security Misconfiguration', 'CORS restricted, no debug mode', 'Addressed'],
            ['API9: Improper Inventory Mgmt', 'All 11 endpoints in Swagger', 'Addressed'],
            ['API10: Unsafe API Consumption', 'No third-party API calls', 'N/A'],
        ],
        caption='Table 11: OWASP API Security Top 10 mapping.'
    )

    add_image_safe(doc, '09_security_results.png', Inches(5.0),
                   'Figure 9: Security testing results, 42/42 tests passed across 8 categories.')

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # SECTION 8
    # ══════════════════════════════════════════════════════════════════
    doc.add_heading('8. Limitations and Challenges', level=1)

    doc.add_heading('8.1 Current Limitations', level=2)
    add_numbered(doc, " All testing is on generated data. Real-world mule patterns are messier, more varied, and more creative.", bold_prefix="Synthetic Data:")
    add_numbered(doc, " We build the transaction graph once at startup from CSV files. Production needs incremental updates.", bold_prefix="Static Graph Analysis:")
    add_numbered(doc, " All data in RAM. Works at current scale but wouldn't work at UPI's actual volume of billions of transactions.", bold_prefix="Everything's in Memory:")
    add_numbered(doc, " Our graph analysis uses hand-crafted features. GNNs could learn subtler patterns.", bold_prefix="No GNNs:")
    add_numbered(doc, " Ensemble weights (25/40/15/10/10) are set manually, not learned from data.", bold_prefix="Fixed Weights:")
    add_numbered(doc, " Isolation Forest trains once in batch. Mule tactics evolve; the model should too.", bold_prefix="No Incremental Learning:")
    add_numbered(doc, " We check direct device sharing only, not multi-hop transitive device chains.", bold_prefix="Single-Hop Device Correlation:")

    doc.add_heading('8.2 Challenges Encountered', level=2)
    add_numbered(doc, " nx.simple_cycles() never returned on dense graphs. Custom DFS with depth cap of 6 solved it.", bold_prefix="Cycle Detection Scalability:")
    add_numbered(doc, " Finding the right signal weights required dozens of experiments.", bold_prefix="Ensemble Calibration:")
    add_numbered(doc, " Evidence items raise privacy questions. Production needs access controls on surfaced information.", bold_prefix="Explainability vs. Privacy:")
    add_numbered(doc, " Developed on Windows, deployed in Linux Docker. File paths and encoding differences.", bold_prefix="Cross-Platform Issues:")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # SECTION 9
    # ══════════════════════════════════════════════════════════════════
    doc.add_heading('9. Roadmap Towards MVP', level=1)

    phases = [
        ('Phase 1: Core Infrastructure (Weeks 1\u20134)', [
            'Real-Time Data Ingestion: Apache Kafka consumer for UPI transaction events.',
            'Persistent Storage: PostgreSQL for historical transactions and audit logs.',
            'Graph Database: Neo4j with native incremental updates.',
            'Redis Cache: Sub-millisecond access for hot data.',
        ]),
        ('Phase 2: Advanced Detection (Weeks 5\u20138)', [
            'Graph Neural Networks: GNN-based node classification with PyTorch Geometric.',
            'Incremental Learning: Online anomaly detection replacing batch Isolation Forest.',
            'Bidirectional Analysis: Incoming patterns with equal sophistication.',
            'Multi-Hop Device Chains: Transitive device sharing across multiple hops.',
        ]),
        ('Phase 3: Scale and Integration (Weeks 9\u201312)', [
            'Kubernetes: Horizontal autoscaling for peak transaction volumes.',
            'UPI Switch Plugin: Inline scoring during payment processing.',
            'Alert Management: Case management workflow with assignment and escalation.',
            'Feedback Loop: Investigator decisions fed back into detection weights and ML model.',
        ]),
        ('Phase 4: Production Hardening (Weeks 13\u201316)', [
            'A/B Testing Framework: Shadow mode for new detection models.',
            'Regulatory Compliance: Automated SAR generation per RBI guidelines.',
            'SOC 2 Certification: Security controls and audit documentation.',
            'Performance Target: Sub-10ms scoring at 10,000 TPS.',
        ]),
    ]
    for title, items in phases:
        doc.add_heading(title, level=2)
        for item in items:
            parts = item.split(': ', 1)
            if len(parts) == 2:
                add_bullet(doc, ' ' + parts[1], bold_prefix=parts[0] + ':')
            else:
                add_bullet(doc, item)

    add_image_safe(doc, '08_roadmap.png', Inches(5.8),
                   'Figure 10: 16-week MVP development roadmap across four phases.')

    doc.add_heading('End-Use Cases', level=2)
    add_numbered(doc, " Dormant accounts suddenly receiving and forwarding large amounts. The system catches the telltale signature.", bold_prefix="Student Mule Accounts:")
    add_numbered(doc, " Dozens of accounts working in concert. Graph analysis maps the entire network.", bold_prefix="Organized Fraud Rings:")
    add_numbered(doc, " Account takeover scenarios. Device fingerprinting catches the new device; behavioral shifts appear in temporal analysis.", bold_prefix="Compromised Legitimate Customers:")
    add_numbered(doc, " Money running in circles to obscure origin. Cycle detection was built for this.", bold_prefix="Layering Operations:")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # SECTION 10
    # ══════════════════════════════════════════════════════════════════
    doc.add_heading('10. Team Composition and Individual Contributions', level=1)
    add_styled_table(doc,
        ['Member', 'Role', 'University', 'DOB', 'Key Contributions'],
        [
            ['[Team Lead Name]', 'Team Leader', '[University]', '[DOB]', 'System architecture, risk engine, ensemble weight tuning, report authoring'],
            ['[Member 2 Name]', 'Backend Dev', '[University]', '[DOB]', 'Graph analysis module, behavioral analysis, DFS cycle detection, BFS chain detection'],
            ['[Member 3 Name]', 'ML Engineer', '[University]', '[DOB]', 'Custom Isolation Forest (NumPy), 17-feature pipeline, Z-score ensemble, SHAP-like explainer'],
            ['[Member 4 Name]', 'Frontend Dev', '[University]', '[DOB]', 'React dashboard (all 8 tabs), network graph vis, API integration, dark theme UX'],
            ['[Member 5 Name]', 'DevOps / QA', '[University]', '[DOB]', 'Docker containerization, security middleware, data generation, temporal analysis module'],
        ],
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # SECTION 11
    # ══════════════════════════════════════════════════════════════════
    doc.add_heading('11. References', level=1)

    references = [
        '[1] NPCI, "UPI Product Statistics," 2024. Available: https://www.npci.org.in/what-we-do/upi/product-statistics',
        '[2] RBI, "Master Direction on Digital Payment Security Controls," RBI/2020-21/74, 2021. Available: https://www.rbi.org.in',
        '[3] NPCI, "UPI Fraud Monitoring and Risk Management Guidelines," 2023. Available: https://www.npci.org.in',
        '[4] S. Panigrahi et al., "A Detailed Study of Rule-Based and Machine Learning Methods for Fraud Detection," J. King Saud Univ. - Comp. Info. Sci., vol. 34, no. 9, pp. 7524\u20137537, 2022.',
        '[5] E. A. Lopez-Rojas et al., "Applying AI and ML in Financial Services," IEEE Access, vol. 10, pp. 76200\u201376215, 2022.',
        '[6] G. Jambhrunkar et al., "MuleTrack: A Lightweight Temporal Learning Framework for Money Mule Detection," IWANN, 2025.',
        '[7] D. Cheng et al., "Graph Neural Networks for Financial Fraud Detection: A Review," arXiv:2411.05815, 2024.',
        '[8] M. Caglayan and S. Bahtiyar, "Money Laundering Detection with Node2Vec," Gazi Univ. J. Sci., vol. 35, no. 3, pp. 854\u2013873, 2022.',
        '[9] Z. Huang, "Enhancing Anti-Money Laundering by Money Mules Detection on Transaction Graphs," ACM GAIB, 2025.',
        '[10] Neo4j Inc., "Accelerate Fraud Detection with Graph Databases," Whitepaper, 2023. Available: https://neo4j.com',
        '[11] F. T. Liu, K. M. Ting, Z.-H. Zhou, "Isolation Forest," 8th IEEE Int. Conf. Data Mining, pp. 413\u2013422, 2008.',
    ]
    for ref in references:
        p = doc.add_paragraph()
        run = p.add_run(ref)
        run.font.name = 'Georgia'
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        p.paragraph_format.space_after = Pt(4)

    add_divider(doc)

    # ── APPENDIX ────────────────────────────────────────────────────
    doc.add_heading('Appendix: Additional Resources', level=1)
    add_styled_table(doc,
        ['Resource', 'Link'],
        [
            ['GitHub Repository', '[Insert GitHub Repo URL]'],
            ['Live Deployed URL', '[Insert Deployed URL]'],
            ['Dashboard Demo Video (YouTube)', '[Insert YouTube Link]'],
            ['API Documentation (Swagger)', '[Insert /docs URL when deployed]'],
            ['Pitch Deck', '[Insert Link]'],
        ],
        header_color="333333"
    )

    add_para(doc, '', size=Pt(8))
    add_divider(doc)
    add_para(doc, 'Submitted by Team FinGuard for the Cyber Security Innovation Challenge (CSIC) 1.0, Stage III Prototype Development.',
             italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x55, 0x55, 0x55))

    # ── SAVE ────────────────────────────────────────────────────────
    doc.save(OUT_PATH)
    print(f"\n[+] Document saved to: {OUT_PATH}")
    print(f"    Size: {os.path.getsize(OUT_PATH) / 1024:.1f} KB")


if __name__ == '__main__':
    print("Generating FINAL_REPORT_CLEANED.docx...")
    build_docx()
    print("Done!")
