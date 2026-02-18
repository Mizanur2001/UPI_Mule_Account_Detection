"""
FinGuard FINAL_REPORT.md → Professional DOCX Generator
=======================================================
Reads the FINAL_REPORT.md and converts it into a beautifully styled
Word document with embedded diagrams, formatted tables, math formulas,
code blocks, and clean typography.

Usage:
    python scripts/generate_fresh_docx.py
"""

import os
import re
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DIAGRAMS = os.path.join(BASE_DIR, "docs", "diagrams")
MD_PATH = os.path.join(BASE_DIR, "FINAL_REPORT.md")
OUT_PATH = os.path.join(BASE_DIR, "FINAL_REPORT.docx")

# ─── Image mapping: markdown image paths → actual filenames ───────────
IMAGE_MAP = {
    "docs/diagrams/01_system_flow.png": "01_system_flow.png",
    "docs/diagrams/02_architecture.png": "02_architecture.png",
    "docs/diagrams/03_graph_patterns.png": "03_graph_patterns.png",
    "docs/diagrams/04_device_flow.png": "04_device_flow.png",
    "docs/diagrams/05_risk_pipeline.png": "05_risk_pipeline.png",
    "docs/diagrams/06_risk_distribution.png": "06_risk_distribution.png",
    "docs/diagrams/07_deployment.png": "07_deployment.png",
    "docs/diagrams/08_roadmap.png": "08_roadmap.png",
    "docs/diagrams/09_security_results.png": "09_security_results.png",
    "docs/diagrams/10_signal_radar.png": "10_signal_radar.png",
}


# ═══════════════════════════════════════════════════════════════════════
# STYLING HELPERS
# ═══════════════════════════════════════════════════════════════════════

def setup_styles(doc):
    """Configure document-wide styles for a classic, professional look."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Georgia'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15

    for level in range(1, 4):
        sname = f'Heading {level}'
        hstyle = doc.styles[sname]
        hstyle.font.name = 'Calibri'
        hstyle.font.color.rgb = RGBColor(0x0D, 0x0D, 0x0D)
        hstyle.font.bold = True
        if level == 1:
            hstyle.font.size = Pt(22)
            hstyle.paragraph_format.space_before = Pt(24)
            hstyle.paragraph_format.space_after = Pt(12)
        elif level == 2:
            hstyle.font.size = Pt(16)
            hstyle.paragraph_format.space_before = Pt(18)
            hstyle.paragraph_format.space_after = Pt(8)
        else:
            hstyle.font.size = Pt(13)
            hstyle.paragraph_format.space_before = Pt(12)
            hstyle.paragraph_format.space_after = Pt(6)


def set_cell_shading(cell, color_hex):
    """Set background shading for a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_table(doc, headers, rows, caption=None, header_color="1A1A1A"):
    """Create a beautifully formatted table."""
    num_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h.strip())
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cell, header_color)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg = "F9F9F9" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            val_str = str(val).strip()
            # Handle bold markers **text**
            if val_str.startswith('**') and val_str.endswith('**'):
                run = p.add_run(val_str.strip('*'))
                run.bold = True
            else:
                # Handle inline bold within cell
                add_rich_runs_to_paragraph(p, val_str)
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.name = 'Georgia'
                if not run.bold:
                    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_shading(cell, bg)

    # Set column widths evenly
    for row in table.rows:
        for cell in row.cells:
            cell.width = Inches(6.5 / num_cols)

    if caption:
        p = doc.add_paragraph()
        run = p.add_run(caption)
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(12)

    return table


def add_image_safe(doc, filename, width=Inches(5.5), caption=None):
    """Add an image if it exists, with a caption."""
    path = os.path.join(DIAGRAMS, filename)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=width)
        if caption:
            cp = doc.add_paragraph()
            run = cp.add_run(caption)
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.space_after = Pt(10)
        return True
    else:
        print(f"  [!] Image not found: {path}")
        return False


def add_rich_runs_to_paragraph(p, text):
    """Parse inline markdown formatting (**bold**, *italic*, `code`) and add runs."""
    # Pattern to find **bold**, *italic*, `code` segments
    parts = re.split(r'(\*\*.*?\*\*|\*[^*]+?\*|`[^`]+?`)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = p.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = p.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x80, 0x00, 0x00)
        else:
            run = p.add_run(part)


def add_para(doc, text, bold=False, italic=False, size=Pt(11), color=None, align=None, space_after=Pt(6)):
    """Add a styled paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Georgia'
    run.font.size = size
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if align:
        p.alignment = align
    p.paragraph_format.space_after = space_after
    return p


def add_rich_paragraph(doc, text, align=None, space_after=Pt(6)):
    """Add a paragraph with inline markdown formatting parsed."""
    p = doc.add_paragraph()
    add_rich_runs_to_paragraph(p, text)
    for run in p.runs:
        if run.font.name is None:
            run.font.name = 'Georgia'
        if run.font.size is None:
            run.font.size = Pt(11)
    if align:
        p.alignment = align
    p.paragraph_format.space_after = space_after
    return p


def add_code_block(doc, code_text):
    """Add a code block with monospace font and gray background."""
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2"/>')
    p._p.get_or_add_pPr().append(shading)
    return p


def add_formula(doc, formula_text):
    """Add a math formula as styled text."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(formula_text)
    run.font.name = 'Cambria Math'
    run.font.size = Pt(11)
    run.italic = True
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    return p


def add_divider(doc):
    """Add a thin horizontal line as a divider."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="CCCCCC"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    return p


# ═══════════════════════════════════════════════════════════════════════
# MARKDOWN PARSER
# ═══════════════════════════════════════════════════════════════════════

def parse_md_table(lines):
    """Parse markdown table lines into headers and rows."""
    if len(lines) < 2:
        return None, None
    
    def parse_row(line):
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        return cells
    
    headers = parse_row(lines[0])
    # lines[1] is the separator (---|---)
    rows = []
    for line in lines[2:]:
        if line.strip():
            rows.append(parse_row(line))
    return headers, rows


def clean_latex(text):
    """Convert LaTeX math notation to readable Unicode for Word."""
    # Remove $$ and $ delimiters
    text = text.strip().strip('$').strip()
    
    # Literal string replacements first (no regex)
    literal_replacements = [
        ('\\sum_{k=1}^{5}', 'Σ(k=1 to 5)'),
        ('\\sum', 'Σ'),
        ('\\cdot', '·'),
        ('\\text{Boost}', 'Boost'),
        ('\\text{base}', 'base'),
        ('\\text{-score}', '-score'),
        ('\\ldots', '…'),
        ('\\in', '∈'),
        ('\\min', 'min'),
        ('\\geq', '≥'),
        ('\\leq', '≤'),
        ('\\times', '×'),
        ('\\{', '{'),
        ('\\}', '}'),
    ]
    for old, new in literal_replacements:
        text = text.replace(old, new)

    # Regex-based replacements for patterns
    regex_replacements = [
        (r'\\frac\{E\[h\(x\)\]\}\{c\(n\)\}', 'E[h(x)] / c(n)'),
        (r'\\frac\{([^}]+)\}\{([^}]+)\}', r'\1 / \2'),
        (r'_\{([^}]+)\}', r'_(\1)'),
        (r'\^\{([^}]+)\}', r'^(\1)'),
    ]
    for pattern, replacement in regex_replacements:
        try:
            text = re.sub(pattern, replacement, text)
        except re.error:
            pass
    
    # Clean up remaining backslashes
    text = text.replace('\\', '')
    return text.strip()


def read_markdown():
    """Read and return the markdown content."""
    with open(MD_PATH, 'r', encoding='utf-8') as f:
        return f.read()


# ═══════════════════════════════════════════════════════════════════════
# BUILD THE DOCUMENT
# ═══════════════════════════════════════════════════════════════════════

def build_docx():
    md_content = read_markdown()
    lines = md_content.split('\n')
    
    doc = Document()
    setup_styles(doc)
    
    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── COVER PAGE ──────────────────────────────────────────────────
    doc.add_paragraph()  # spacer
    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('CYBER SECURITY INNOVATION CHALLENGE 1.0')
    run.font.name = 'Calibri'
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = RGBColor(0x0D, 0x0D, 0x0D)

    add_para(doc, '', size=Pt(6))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('PROTOTYPE DEVELOPMENT: Stage III')
    run.font.name = 'Calibri'
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    add_divider(doc)

    for line in [
        'Problem Statement Domain: Mule Accounts & Collusive Fraud in UPI',
        'Problem Statement: Detection of Mule Account Networks and Coordinated Fund Laundering in UPI Payment Ecosystems',
    ]:
        add_para(doc, line, size=Pt(11), align=WD_ALIGN_PARAGRAPH.CENTER,
                 color=RGBColor(0x44, 0x44, 0x44))

    add_para(doc, '', size=Pt(12))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('FinGuard')
    run.font.name = 'Calibri'
    run.font.size = Pt(36)
    run.bold = True
    run.font.color.rgb = RGBColor(0x0D, 0x0D, 0x0D)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Real-Time Multi-Signal Mule Account Detection System')
    run.font.name = 'Georgia'
    run.font.size = Pt(14)
    run.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    add_para(doc, '', size=Pt(24))

    # Project Links table on cover
    add_styled_table(doc,
        ['Resource', 'Link'],
        [
            ['GitHub Repository', '[Insert GitHub Repo URL]'],
            ['Live Deployed URL', '[Insert Deployed URL]'],
            ['Dashboard Demo Video (YouTube)', '[Insert YouTube Link]'],
            ['Pitch Deck / Presentation', '[Insert Link if applicable]'],
        ],
        header_color="333333"
    )

    add_para(doc, '', size=Pt(12))

    # Team table on cover
    add_styled_table(doc,
        ['Role', 'Name', 'Institute', 'Enrolment No.', 'Email ID'],
        [
            ['Team Lead', '[Name]', '[Institute]', '[Enrolment No.]', '[Email]'],
            ['Member 2', '[Name]', '[Institute]', '[Enrolment No.]', '[Email]'],
            ['Member 3', '[Name]', '[Institute]', '[Enrolment No.]', '[Email]'],
            ['Member 4', '[Name]', '[Institute]', '[Enrolment No.]', '[Email]'],
            ['Member 5', '[Name]', '[Institute]', '[Enrolment No.]', '[Email]'],
        ],
        header_color="333333"
    )

    doc.add_page_break()

    # ── TABLE OF CONTENTS ───────────────────────────────────────────
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. Project Overview and Team Details',
        '2. Problem Statement and Background',
        '3. Literature Review / Existing Solutions',
        '4. Proposed Solution and Technical Architecture',
        '5. Innovation and Novelty Elements',
        '6. Unique Selling Proposition (USP), Business Model, and Industry Relevance',
        '7. Prototype Demonstration, Security Testing, and Deployment Details',
        '8. Limitations and Challenges',
        '9. Roadmap Towards MVP',
        '10. Team Composition and Individual Contributions',
        '11. References',
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.name = 'Georgia'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        p.paragraph_format.space_after = Pt(3)

    doc.add_page_break()

    # ── PARSE AND RENDER MARKDOWN BODY ──────────────────────────────
    # We parse the markdown line by line, handling:
    # - Headings (##, ###, ####)
    # - Tables (| ... | ... |)
    # - Images (![alt](path))
    # - Code blocks (``` ... ```)
    # - Math blocks ($$ ... $$)
    # - Ordered lists (1. ...)
    # - Unordered lists (- ...)
    # - Block quotes (> ...)
    # - Paragraphs (regular text)
    # - Horizontal rules (---)
    # - Figure captions (*Figure N: ...*)
    
    i = 0
    # Skip past the cover content in markdown (until "## 1.")
    # Find the line that starts Section 1
    while i < len(lines):
        if lines[i].strip().startswith('## 1. Project Overview'):
            break
        i += 1

    in_code_block = False
    code_block_lines = []
    in_math_block = False
    math_block_lines = []
    table_lines = []
    table_caption_next = False
    skip_ascii_art = False
    
    section_count = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # ── Code block handling ───────────────────────────────────
        if stripped.startswith('```'):
            if in_code_block:
                # End code block
                code_text = '\n'.join(code_block_lines)
                # Check if it's an ASCII art diagram (architecture box)
                if '┌' in code_text or '├' in code_text or '└' in code_text:
                    add_code_block(doc, code_text)
                else:
                    add_code_block(doc, code_text)
                code_block_lines = []
                in_code_block = False
            else:
                # Start code block
                in_code_block = True
                code_block_lines = []
            i += 1
            continue

        if in_code_block:
            code_block_lines.append(line)
            i += 1
            continue

        # ── Math block handling ($$...$$) ─────────────────────────
        if stripped.startswith('$$') and not stripped.endswith('$$'):
            in_math_block = True
            math_block_lines = []
            i += 1
            continue
        
        if in_math_block:
            if stripped.startswith('$$') or stripped.endswith('$$'):
                formula = clean_latex(' '.join(math_block_lines))
                add_formula(doc, formula)
                in_math_block = False
                math_block_lines = []
            else:
                math_block_lines.append(stripped)
            i += 1
            continue

        # Single-line math block: $$formula$$
        if stripped.startswith('$$') and stripped.endswith('$$') and len(stripped) > 4:
            formula = clean_latex(stripped[2:-2])
            add_formula(doc, formula)
            i += 1
            continue

        # ── Table handling ────────────────────────────────────────
        if stripped.startswith('|') and '|' in stripped[1:]:
            table_lines.append(stripped)
            i += 1
            continue
        elif table_lines:
            # Flush table
            _flush_table(doc, table_lines)
            table_lines = []
            # Check if current line is a table caption
            if stripped.startswith('Table ') and ':' in stripped:
                p = doc.add_paragraph()
                run = p.add_run(stripped)
                run.italic = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(12)
                i += 1
                continue
            # Don't increment i, process this line normally
            continue

        # ── Empty line ────────────────────────────────────────────
        if not stripped:
            i += 1
            continue

        # ── Horizontal rule ───────────────────────────────────────
        if stripped == '---' or stripped == '***':
            add_divider(doc)
            i += 1
            continue

        # ── Headings ──────────────────────────────────────────────
        if stripped.startswith('#### '):
            heading_text = stripped[5:].strip()
            doc.add_heading(heading_text, level=3)
            i += 1
            continue

        if stripped.startswith('### '):
            heading_text = stripped[4:].strip()
            doc.add_heading(heading_text, level=2)
            i += 1
            continue

        if stripped.startswith('## '):
            heading_text = stripped[3:].strip()
            # Add page break before major sections (except the first)
            if section_count > 0:
                doc.add_page_break()
            section_count += 1
            doc.add_heading(heading_text, level=1)
            i += 1
            continue

        # ── Images ────────────────────────────────────────────────
        img_match = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', stripped)
        if img_match:
            alt_text = img_match.group(1)
            img_path = img_match.group(2)
            # Look up the image filename
            filename = IMAGE_MAP.get(img_path, os.path.basename(img_path))
            add_image_safe(doc, filename, Inches(5.5))
            i += 1
            continue

        # ── Figure captions (*Figure N: ...*) ─────────────────────
        if stripped.startswith('*Figure') and stripped.endswith('*'):
            caption_text = stripped.strip('*').strip()
            p = doc.add_paragraph()
            run = p.add_run(caption_text)
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(10)
            i += 1
            continue

        # ── Block quotes (> ...) ──────────────────────────────────
        if stripped.startswith('> '):
            quote_text = stripped[2:].strip()
            # Clean markdown bold/italic
            quote_text = re.sub(r'\*\*\[', '[', quote_text)
            quote_text = re.sub(r'\]\*\*', ']', quote_text)
            p = doc.add_paragraph()
            p.style = 'Quote' if 'Quote' in [s.name for s in doc.styles] else 'Normal'
            run = p.add_run(quote_text)
            run.font.name = 'Georgia'
            run.font.size = Pt(10)
            run.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            p.paragraph_format.left_indent = Cm(1.5)
            p.paragraph_format.space_after = Pt(8)
            i += 1
            continue

        # ── Ordered lists (1. ...) ────────────────────────────────
        ol_match = re.match(r'^(\d+)\.\s+(.*)', stripped)
        if ol_match:
            item_text = ol_match.group(2)
            p = doc.add_paragraph(style='List Number')
            p.text = ''
            add_rich_runs_to_paragraph(p, item_text)
            for run in p.runs:
                if run.font.name is None:
                    run.font.name = 'Georgia'
                if run.font.size is None:
                    run.font.size = Pt(11)
            i += 1
            continue

        # ── Unordered lists (- ...) ───────────────────────────────
        if stripped.startswith('- '):
            item_text = stripped[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            p.text = ''
            add_rich_runs_to_paragraph(p, item_text)
            for run in p.runs:
                if run.font.name is None:
                    run.font.name = 'Georgia'
                if run.font.size is None:
                    run.font.size = Pt(11)
            i += 1
            continue

        # ── Inline math ($...$) in a line ────────────────────────
        # Handle lines with inline math by converting to readable text
        if '$' in stripped:
            cleaned = re.sub(r'\$\$([^$]+)\$\$', lambda m: clean_latex(m.group(1)), stripped)
            cleaned = re.sub(r'\$([^$]+)\$', lambda m: clean_latex(m.group(1)), cleaned)
            add_rich_paragraph(doc, cleaned)
            i += 1
            continue

        # ── Regular paragraphs ────────────────────────────────────
        add_rich_paragraph(doc, stripped)
        i += 1

    # Flush any remaining table
    if table_lines:
        _flush_table(doc, table_lines)

    # ── FINAL FOOTER ────────────────────────────────────────────────
    add_divider(doc)
    add_para(doc, 'Submitted by Team FinGuard for the Cyber Security Innovation Challenge (CSIC) 1.0, Stage III Prototype Development.',
             italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x55, 0x55, 0x55))

    # ── SAVE ────────────────────────────────────────────────────────
    doc.save(OUT_PATH)
    print(f"\n[+] Document saved to: {OUT_PATH}")
    print(f"    Size: {os.path.getsize(OUT_PATH) / 1024:.1f} KB")


def _flush_table(doc, table_lines):
    """Parse and render a markdown table."""
    if len(table_lines) < 3:
        # Not enough lines for a proper table
        for line in table_lines:
            add_rich_paragraph(doc, line)
        return

    # Filter out separator rows (---|---)
    header_line = table_lines[0]
    data_lines = []
    for j, line in enumerate(table_lines):
        if j == 0:
            continue
        # Check if separator row
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        if all(re.match(r'^:?-+:?$', c) for c in cells if c):
            continue
        data_lines.append(line)

    headers = [c.strip() for c in header_line.strip().strip('|').split('|')]
    headers = [h for h in headers if h]  # Remove empty strings

    rows = []
    for line in data_lines:
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        cells = [c for c in cells if c is not None]
        # Ensure same number of columns
        while len(cells) < len(headers):
            cells.append('')
        cells = cells[:len(headers)]
        rows.append(cells)

    if headers and rows:
        add_styled_table(doc, headers, rows)


# ═══════════════════════════════════════════════════════════════════════
# ENTRY POINT
# ═══════════════════════════════════════════════════════════════════════

if __name__ == '__main__':
    print("=" * 60)
    print("  FinGuard FINAL_REPORT.md → DOCX Generator")
    print("=" * 60)
    print(f"\n  Source:  {MD_PATH}")
    print(f"  Output:  {OUT_PATH}")
    print(f"  Diagrams: {DIAGRAMS}")
    print()
    
    if not os.path.exists(MD_PATH):
        print(f"[ERROR] Markdown file not found: {MD_PATH}")
        exit(1)
    
    print("Generating FINAL_REPORT.docx from FINAL_REPORT.md...")
    build_docx()
    print("\nDone!")
